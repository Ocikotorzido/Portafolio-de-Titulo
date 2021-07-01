import csv
import smtplib
import datetime
import os, os.path
from threading import Thread

import docx
import pandas
import docx2pdf
#import rut_chile
from .models import *
from django.apps import apps
from django.http import HttpResponse
from django.http import FileResponse
from django.utils.encoding import smart_str
from django.contrib.auth.models import User
from django.shortcuts import render, redirect
from django.template.defaultfilters import date
from django.contrib.auth.forms import AuthenticationForm
from django.utils.translation import get_language, activate
from django.contrib.auth import login as iniciarSesion, logout, authenticate

from Taller.settings import EMAIL_HOST_USER, EMAIL_HOST_PASSWORD

def to_index(request):
    """Redirección hacia index"""
    return redirect('index')

def login(request):
    """Módulo que permite al usuario registrarse e/o iniciar sesión."""    

    # Si el usuario ya está loggeado,
    # Se redirecciona al index.
    if request.user.is_authenticated:
        return redirect('index')

    if request.method == 'POST':
        # Se rescatan los campos de usuario y contraseña.
        only_numbers = lambda texto: ''.join([numero if numero in ['0','1','2','3','4','5','6','7','8','9'] else '' for numero in texto])
        usuario = only_numbers(request.POST['username'])
        contra = request.POST['password']

        # Se verifica que las crenciales sean válidas.
        usuarioLogeado = authenticate(username = usuario, password = contra)

        # Debiera devolver una instancia de sesión.
        # De lo contrario, devuelve 'None'.
        if usuarioLogeado is not None:

            # La función login() iniciará la sesión.
            iniciarSesion(request, usuarioLogeado)

            # Si la autenticación es correcta, 
            # entonces la plantilla se renderizará
            # como un usuario válidamente autenticado.
            return redirect('index')
        
        # Si la autenticación falla, 
        # la sesión NO existirá.
        else:
            # Se verifica que exista el nombre de usuario en la base de datos.
            if len(User.objects.filter(username=usuario)):
                # Si existe, se le informa al usuario que la contraseña es incorrecta.
                return render(request, 'mantenedor/login.html', {'errores':'CONTRASEÑA_INCORRECTA'})
            else:
                # Si no existe, se le informa al usuario que el nombre de usuario no existe.
                return render(request, 'mantenedor/login.html', {'errores':'NO_EXISTE_NOMBRE_USUARIO'})

    # Página de inicio de sesión.
    return render(request, 'mantenedor/login.html')

def registro (request):
    formulario = None
    if request.method == 'POST':
        formulario = FormularioRegistro(data = request.POST)
        # try:
        #     if not rut_chile.is_valid_rut(formulario['username']):
        #         return render(request, 'Mantenedor/registro.html', {formulario:formulario, 'errores':'rut mal ingresado'})
        # except ValueError:
        #     return render(request, 'Mantenedor/registro.html', {formulario:formulario, 'errores':'rut mal ingresado'})
        if formulario.is_valid():
            usuario_guardado = formulario.save()
            if usuario_guardado is not None:
                iniciarSesion(request,usuario_guardado)
                context = {'perfil':
                            {'nivel':'admin'},
                            'level':'admin'
                            }
                return render(request, 'TEMPORAL/perfil.html', context)                
    else:
        formulario = FormularioRegistro()

    context = {'formulario': formulario}
    return render(request, 'TEMPORAL/registro.html', context)

def perfil(request):
    return render(
        request,
        'TEMPORAL/perfil.html'
    )

def salir(request):
    logout(request)
    return redirect('index')

def obtener_usuario(request):
    nivel = None
    if request.user.is_authenticated:
        try:
            nivel = Perfil.objects.filter(id_auth_user = request.user.id)[0].nivel
        except IndexError:
            logout(request)
            nivel = 'ERROR'
    return {'perfil':{'nivel':nivel}}

def index (request):
    context = obtener_usuario(request)
    return render (request, 'mantenedor/index.html',context)


def registro_cliente (request):
    formulario = FormularioRegistro()
    context = dict()
    permitir_sesion = False
    if request.method == 'POST':
        formulario = FormularioRegistro(data = request.POST)
        if formulario.is_valid():
            usuario_guardado = formulario.save()
            if usuario_guardado is not None:
                permitir_sesion = True
        else:
            context['formulario'] = formulario
            return render (request, 'mantenedor/registro_cliente.html',context)

        dato_or_zero = lambda dato_crudo: 0 if not dato_crudo else dato_crudo
        only_numbers = lambda texto: ''.join([numero if numero in ['0','1','2','3','4','5','6','7','8','9'] else '' for numero in texto])

        mi_rut = only_numbers(formulario.cleaned_data['username'])
        mi_nombre = request.POST['first_name']
        mi_apellido = request.POST['last_name']
        mi_direccion = request.POST['direccion']
        mi_telefono = dato_or_zero(request.POST['telefono'])
        mi_celular = dato_or_zero(request.POST['celular'])
        mi_email = request.POST['email']

        id_user_auth = User.objects.get(username=formulario.cleaned_data['username']).id
        id_cliente = Cliente.objects.count()+1
        cliente = Cliente(id_cliente,mi_nombre,mi_apellido,mi_direccion,
                            mi_telefono,mi_celular,mi_email,mi_rut)

        n_perfil = Perfil.objects.all().count()+1
        perfil = Perfil(n_perfil,id_user_auth, id_cliente,'CLIENTE')
        perfil.save()

        cliente.save()
        if permitir_sesion:
            iniciarSesion(request,usuario_guardado)
            nivel = Perfil.objects.filter(id_auth_user = request.user.id)[0].nivel
        return render(request, 'Mantenedor/index.html', {'perfil':{'nivel':nivel}})
    context['formulario'] = formulario
    return render (request, 'mantenedor/registro_cliente.html',context)


def ver_perfil(request):

    id_cliente = Perfil.objects.filter(id_auth_user = request.user.id)[0].id_usuario
    rut = User.objects.filter(id = request.user.id)[0].username

    cliente = Cliente.objects.get(id_cliente=id_cliente)
    context = {'cliente':cliente}

    if request.method == 'POST':
        only_numbers = lambda texto: ''.join([numero if numero in ['0','1','2','3','4','5','6','7','8','9'] else '' for numero in texto])
        
        mi_nombre = request.POST['nombre']
        mi_apellido = request.POST['apellido']
        mi_direccion = request.POST['direccion']
        mi_telefono = only_numbers(request.POST['telefono'])
        mi_celular = only_numbers(request.POST['celular'])
        mi_email = request.POST['email']

        cliente = Cliente()

        cliente.id_cliente
        
        cliente.nombre = mi_nombre
        cliente.apellido = mi_apellido
        cliente.direccion = mi_direccion
        cliente.telefono = mi_telefono
        cliente.celular = mi_celular
        cliente.email = mi_celular
        cliente.email = mi_email

        cliente = Cliente(id_cliente,mi_nombre,mi_apellido,mi_direccion,mi_telefono,
                            mi_celular,mi_email,rut)
        cliente.save()

        a=User.objects.get(username=rut)
        a.first_name = mi_nombre
        a.last_name = mi_apellido
        a.email = mi_email
        a.save()
        
        context = {'cliente':cliente}

        return render(request,'mantenedor/ver_perfil.html',context)

    return render(request,'mantenedor/ver_perfil.html',context)





def servicios (request):
    servicios = TipoServicio.objects.all()
    context = {'servicios': servicios}
    
    return render (request, 'mantenedor/servicios.html',context)

def reservas (request):
    cliente = dict()
    cliente['nombre'] = f'{request.user.first_name} {request.user.last_name}'
    cliente['correo'] = f'{request.user.email}'
    vehiculo = dict()
    solicitud = ''
    servicios_disponibles = {
        'servicio_1': 'Cambio de neumáticos',
        'servicio_2': 'Cambio de aceite',
        'servicio_3': 'Cambio de pastillas',
        'servicio_4': 'Cambio de correas',
        'servicio_5': 'cambio de bujías',
        'servicio_6': 'Cambio de amortiguadores',
        'servicio_7': 'Cambio de baterías',
        'servicio_8': 'Cambio de filtros',
    }
    servicios = TipoServicio.objects.all()
    context = {'servicios': servicios }
    if request.method == 'POST':
        context['method'] = 'POST'
        vehiculo['marca'] = request.POST['marca']
        vehiculo['modelo'] = request.POST['modelo']
        vehiculo['year'] = request.POST['year']
        
        id_cliente = Perfil.objects.filter(id_auth_user = request.user.id)[0].id_usuario
        id_reserva = Reservas.objects.count()+1
        reservas = Reservas(id_reserva, 
                            id_cliente, 
                            vehiculo['marca'], 
                            vehiculo['modelo'], 
                            vehiculo['year'], 
                            datetime.datetime.now(), 
                            0)
        reservas.save()
        
        services = list()
        for service in request.POST:
            if 'servicio' in service:
                services.append(servicios_disponibles[service])
                id_servicio = service.replace('servicio_','')
                id_detalle = DetalleSer.objects.count()+1
                detalle = DetalleSer(id_detalle, id_reserva, id_servicio)
                detalle.save()
        solicitud = '\n'.join(services)
        
        tupla_datos = (cliente,vehiculo, solicitud)
        al_recepcionista = Thread(target=enviar_correo, args=(tupla_datos,))
        al_recepcionista.start()
        return render (request, 'mantenedor/reservas.html', context)
    return render (request, 'mantenedor/reservas.html', context)

def modificar_reserva(request, id_reserva, id_mecanico, confirmacion):
    if confirmacion not in [0,1]:
        return HttpResponse(status=403)
    reserva = Reservas.objects.get(id_reserva=id_reserva)
    reserva.confirmacion = confirmacion
    reserva.save()
    ot = Ot.objects.filter(reservas_id_reserva_id=id_reserva)
    if ot.count() == 0:
        id_ot = Ot.objects.count() + 1
        ot = Ot( id_ot, id_mecanico, id_reserva, datetime.datetime.now() )
        ot.save()
    else:
        ot = Ot.objects.get(reservas_id_reserva=id_reserva)
        ot.delete()
    return HttpResponse(status=200)

def ver_reservas (request):
    reservas = Reservas.objects.all()
    mecanicos = Empleado.objects.filter(cargo_id_tipo_cargo=3)
    context = {'reservas': reservas,
               'mecanicos': mecanicos }
    return render (request, 'mantenedor/ver_reservas.html', context)


def orden_trabajo (request):
    context = dict()


    
    ot = Ot.objects.all()
    context['ot'] = ot

    
    detalles = DetalleSer.objects.all()
    context['detalles'] = detalles
    
    reservas = Reservas.objects.filter(confirmacion='1')
    context['reservas'] = reservas
    
    servicios = TipoServicio.objects.all()
    context['servicios'] = servicios
    
    return render (request, 'mantenedor/orden_trabajo.html', context)

def orden_pedido (request):
    proveedores = Proveedor.objects.all()
    autos = InfoAuto.objects.all().order_by('-id_informe')
    context = {'proveedores':proveedores,
               'autos': autos }
    if request.method == 'POST':
         #= request.POST['']
        auto = request.POST['auto']
         
        producto = request.POST['producto']
        cantidad = request.POST['cantidad']
        fecha_pedido = datetime.datetime.now()
        hora_pedido = datetime.datetime.now().strftime('%H_%M')
        fecha_entrega = datetime.datetime.now() + datetime.timedelta(days=2)
        proveedor = request.POST['proveedor']
        
        id_pedido = Op.objects.count()+1
        op = Op(id_pedido, producto, cantidad, fecha_pedido, hora_pedido, fecha_entrega, proveedor)
        op.save()
        
        id_detalle = DetallePedido.objects.count()+1
        dp = DetallePedido(id_detalle, id_pedido, proveedor)
        dp.save()
        
    return render (request, 'mantenedor/orden_pedido.html', context)

def registrar_proveedor(request):
    nivel = None
    if request.user.is_authenticated:
        try:
            nivel = Perfil.objects.filter(id_auth_user = request.user.id)[0].nivel
        except IndexError:
            logout(request)
            nivel = 'ERROR'
            
    context={'perfil':{'nivel':nivel}}
    
    if request.method == 'POST':
        only_numbers = lambda texto: ''.join([numero if numero in ['0','1','2','3','4','5','6','7','8','9'] else '' for numero in texto])
        
        mi_rut = only_numbers(request.POST['rut'])
        mi_nombre = request.POST['nombre']
        mi_telefono = only_numbers(request.POST['telefono'])
        mi_email = request.POST['email']
        mi_rubro = request.POST['rubro']

        proveedor = Proveedor()
        id_proveedor = Proveedor.objects.count()+1

        proveedor.rut = mi_rut
        proveedor.telefono = mi_telefono
        proveedor.email = mi_email
        proveedor.nombre = mi_nombre
        proveedor.rubro = mi_rubro

        proveedor = Proveedor(id_proveedor,
                              mi_rut,
                              mi_nombre,
                              mi_telefono,
                              mi_email,
                              mi_rubro)

        proveedor.save()
        context['msj'] = '¡Registrado exitosamente!'
        return render (request, 'mantenedor/registro_proveedor.html', context)

    return render (request, 'mantenedor/registro_proveedor.html', context)

def ver_proveedores (request):
    
    proveedores = Proveedor.objects.all()
    context = {'proveedores': proveedores}

    return render (request, 'mantenedor/ver_proveedores.html', context)




def registro_vehiculo(request):
    if request.method == 'POST':
        

        mi_cliente = request.POST['nombre']
        mi_rut = request.POST['rut']
        mi_direccion = request.POST['direccion']
        mi_contacto = request.POST['contacto']
        mi_fecha = request.POST["fecha"]
        mi_modelo    = request.POST['modelo']
        mi_marca = request.POST['marca']
        mi_patente = request.POST['patente']
        mi_color = request.POST['color']
        mi_year = request.POST['year']
        mi_kilometraje = request.POST['kilometraje']
        mi_combustible = request.POST['combustible']
        mi_aceite = request.POST['aceite']
        mi_refrigeracion = request.POST['refrigeracion']
        mi_frenos = request.POST['frenos']
        mi_alta_der = request.POST['alta_der']
        mi_alta_izq = request.POST['alta_izq']
        mi_baja_der = request.POST['baja_der']
        mi_baja_izq = request.POST['baja_izq']
        mi_intermitente = request.POST['intermitente']
        mi_observaciones = request.POST['observaciones']

        vehiculo = InfoAuto()
        id_informe = InfoAuto.objects.count()+1

        #asignacion de datos

        vehiculo.cliente = mi_cliente
        vehiculo.rut = mi_rut
        vehiculo.direccion = mi_direccion
        vehiculo.contacto = mi_contacto
        vehiculo.fecha = mi_fecha
        vehiculo.modelo = mi_modelo
        vehiculo.marca = mi_marca
        vehiculo.patente = mi_patente
        vehiculo.color = mi_color
        vehiculo.year = mi_year
        vehiculo.kilometraje = mi_kilometraje
        vehiculo.combustible = mi_combustible
        vehiculo.a_motor = mi_aceite
        vehiculo.refrigerante = mi_refrigeracion
        vehiculo.liq_frenos = mi_frenos
        vehiculo.a_der = mi_alta_der
        vehiculo.a_izq = mi_alta_izq
        vehiculo.b_der = mi_baja_der
        vehiculo.b_izq = mi_baja_izq
        vehiculo.intermitente = mi_intermitente
        vehiculo.observaciones = mi_observaciones

        # Búsqueda del número de identidad.
        
        
        id_emple =    Perfil.objects.filter(id_auth_user = request.user.id)[0].id_usuario
        
        #envio de datos
        nuevo_vehiculo = InfoAuto(id_informe,mi_cliente,mi_rut,mi_direccion,mi_contacto,mi_fecha,mi_modelo,
                                    mi_marca,mi_patente,mi_color,mi_year,mi_kilometraje,
                                    mi_combustible,mi_aceite,mi_refrigeracion,mi_frenos,
                                    mi_alta_der,mi_alta_izq,mi_baja_der,mi_baja_izq,mi_intermitente,mi_observaciones,
                                    id_emple)
        nuevo_vehiculo.save()
        return render(request,'mantenedor/registro_vehiculo.html')
    
    return render(request, 'mantenedor/registro_vehiculo.html')


     


def presupuesto(request):
    return render (request, 'mantenedor/presupuesto.html')


def faq(request):
    return render (request, 'mantenedor/faq.html')
           

def empleado (request):
    cargos = Cargo.objects.all()
    context = {'cargos': cargos }
    return render(request,'mantenedor/registro_empleado.html',context)

def exportar(request):
    nivel = None
    if request.user.is_authenticated:
        nivel = Perfil.objects.filter(id_auth_user = request.user.id)[0].nivel
    context={'perfil':{'nivel':nivel}}
    return render(request,'mantenedor/exportar.html',context)

def agregar_empleado(request):
    if request.method == 'POST':
        mi_cargo = request.POST['cargo']
        mi_rut = request.POST['rut']
        mi_nombre = request.POST['nombre']
        mi_apellido = request.POST['apellido']
        mi_contacto = request.POST['contacto']
        mi_password = request.POST['password']

        empleado = Empleado()
        id_empleado = Empleado.objects.count()+1
        empleado.rut = mi_rut
        empleado.nombre = mi_nombre
        empleado.apellido = mi_apellido
        empleado.contacto = mi_contacto
        
        # Creación del empleado en la tabla 'auth_user'.
        nuevo_empleado = User.objects.create_user(mi_rut,mi_contacto,mi_password) 
        nuevo_empleado.save()

        # Búsqueda del número de identidad.
        id_user_auth = User.objects.get(username=mi_rut).id

        # Creación de empleado en la tabla 'perfil'.
        n_perfil = Perfil.objects.all().count()+1
        perfil = Perfil(n_perfil,id_user_auth, id_empleado, 'EMPLEADO')
        perfil.save()
        
        # Creación de un nuevo empleado.
        empleado = Empleado(id_empleado,
                                mi_rut,
                                mi_nombre,
                                mi_apellido,
                                mi_contacto,
                                mi_cargo)
        empleado.save()
        return render(request, 'mantenedor/registro_empleado.html', {'mensaje':'Empleado_registrado'})

def generar_informe(request, informe_de, parametros, tipo):
    """Generar informes sobre algun/a persona/objeto.
    informe_de -> Tipo de informe, 
        puede ser: empleado, cliente, 
        administrador, vehículo, etc.
    parametros -> Valores/restricciones del informe,
        puede ser: Todo, último mes,
        última semana, último año,
        o incluso descartados.
    tipo -> formato de salida del informe,
        puede ser: excel, pdf, csv, word.
        
        El flujo de datos va de CSV a XLSX, 
        luego pasa a Docx y finalmente se 
        convierte en PDF, previo es obligado.
    """
    
    tablas_db = apps.all_models['Mantenedor']
    informes = {
        # 'NOMBRE_EN_FORM_DE_HTML' : tablas_db['NOMBRE_TABLA_ORACLE'],
        'empleado': tablas_db['empleado'],
        'cliente': tablas_db['cliente'],
        'proveedor': tablas_db['proveedor'],
        'administrador': tablas_db['perfil'],
        #'vehiculo': tablas_db['reserva'],
    }
    
    
    # Abreviación, extensión y 'content-type' de archivos y sus formatos.
    tipos_admitidos = {
        'csv': 
            ['csv', 'text/csv'],
        'excel': 
            ['xlsx','application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'],
        'word': 
            ['docx',' application/vnd.openxmlformats-officedocument.wordprocessingml.document'],
        'pdf': 
            ['pdf', 'application/pdf'],
        }
    
    # Se activa la traducción de fechas a español.
    activate('es')
    
    # Obtenemos el nombre del mes en español.
    today = datetime.date.today()
    mes = date(today, 'F')
    
    # Se asigna la fecha actual de dos maneras,
    # Una para ser escrita en los nombres de archivos,
    # La otra para ser escrita dentro de los informes.
    now = datetime.datetime.now().strftime('%Y-%m-%d__%H_%M_%S')
    now_ = datetime.datetime.now().strftime(f'%d de {mes} de %Y, %H:%M %p')
                                            
    # Se normaliza el dato.
    tipo = tipo.lower()
    
    # Validación de formato.
    if tipo not in ['csv', 'excel', 'word', 'pdf']:
        return HttpResponse('ERROR, el tipo de formato no es válido!')
    
    # Se define el nombre del archivo.
    nombre_archivo = f'informe_{informe_de}_{now}'
    nombre_archivo_con_extension = f'informe_{informe_de}_{now}.{tipos_admitidos[tipo][0]}'
    
    # Se define el tipo de respuesta y la cabecera.
    response = HttpResponse(
        content_type=f'{tipos_admitidos[tipo][1]}',
        headers={'Content-Disposition': 
            f'attachment; filename="{nombre_archivo_con_extension}"'},
    )
    
    # Obtener los títulos de una tabla.
    fields = informes[informe_de]._meta.get_fields()
    titulos = list()
    for titulo in fields:
        titulos.append(titulo.name)
    
    writer = csv.writer(response)
    writer.writerow(titulos)
    
    # Se obtiene la info de la base de datos.
    nombre_campos = informes[informe_de]._meta.get_fields()
    for fila in informes[informe_de].objects.all():
        temp = list()
        for columna in nombre_campos:
            temp.append(fila.serializable_value(columna.name))
            #if informe_de == 'administrador':
                #if fila.nivel.lower() in ['admin', 'administrador']:
                #    temp.append(fila.serializable_value(columna.name))
            #else:
                #temp.append(fila.serializable_value(columna.name))
        writer.writerow(temp)
    
    # Devuelve un archivo CSV.
    if tipo == 'csv': 
        return response
    
    # Se define la ubicación de los archivos temporales.
    temp_folder = f'{os.path.realpath(".")}\\__temp\\'
    temp_csv = f'{temp_folder}__temp.csv'
    
    # Se corrobora que exista la carpeta temporal.
    if not os.path.exists(temp_folder):
        os.mkdir(temp_folder)
    
    # Se escribe el CSV en físico.
    temp = open(temp_csv, 'wb')
    temp.write(response.content)
    temp.close()
    
    # Devuelve un archivo XLSX.
    if tipo == 'excel': 
        # Pandas lee el CSV desde un archivo.
        archivo_leido = pandas.read_csv(temp_csv)
        
        # Se convierte a excel y se almacena como archivo XLSX.
        archivo_leido.to_excel(f'{temp_folder}{nombre_archivo_con_extension}', 
                                index = None, header=True, sheet_name=f'{informe_de}')
        
        # Devuelve un archivo XLSX.
        return FileResponse(open(f'{temp_folder}{nombre_archivo_con_extension}', 'rb'))
    
    # Se crea y se rellena un archivo DOCX.
    document = docx.Document()
    document.add_heading(f'Informe de {informe_de}', 0)
    document.add_paragraph(f'Con fecha {now_}.')
    
    with open(temp_csv, newline='') as f:
        csv_reader = csv.reader(f)
        csv_headers = next(csv_reader)
        csv_cols = len(csv_headers)
        table = document.add_table(rows=2, cols=csv_cols)
        hdr_cells = table.rows[0].cells
        for i in range(csv_cols):
            hdr_cells[i].text = csv_headers[i]

        for row in csv_reader:
            row_cells = table.add_row().cells
            for i in range(csv_cols):
                row_cells[i].text = row[i]
    document.add_page_break()    
    document.save(f'{temp_folder}{nombre_archivo}.docx')

    # Devuelve un archivo DOCX.
    if tipo == 'word': 
        return FileResponse(open(f'{temp_folder}{nombre_archivo}.docx', 'rb'))
    
    if tipo == 'pdf':
        try:
            # Solución 1.
            # Usando MS-Office 365
            print('\nUsando Office 365\n')
            docx2pdf.convert(f'__temp\{nombre_archivo}.docx')
        except:
            import subprocess
            # Solución 2.
            # Usando LibreOffice.
            print('\nUsando LibreOffice\n')
            path_to_soffice_exe = '"C:\Program Files\LibreOffice\program\soffice.exe"'
            to_pdf = '-headless -convert-to pdf'
            outdir = '-outdir .\__temp'
            res = subprocess.run(f'{path_to_soffice_exe} {to_pdf} {outdir} "__temp\{nombre_archivo}.docx"')
            print(f'\n\n{res}\n\n')
        return FileResponse(open(f'{temp_folder}{nombre_archivo}.pdf', 'rb'))
    else:
        return HttpResponse('Error con el servidor...')

def enviar_correo(tupla_datos):
    cliente,vehiculo,solicitud = tupla_datos[0],tupla_datos[1],tupla_datos[2]
    hoy = datetime.datetime.now().strftime(f'%d de %m de %Y, a las %H:%M %p')
    to = [EMAIL_HOST_USER]
    subject = f'Nueva solicitud de {cliente.get("nombre")} para reserva & presupuesto, {hoy}'
    body = f"""<b>Datos del cliente</b>\nEl cliente {cliente.get('nombre')} solicita presupuesto el dia {hoy}\nResponder al correo: {cliente.get('correo')}\n\n"""
    body += f"""<b>Datos del vehículo</b>\nMarca: {vehiculo.get('marca')}\nModelo: {vehiculo.get('modelo')}\nAño: {vehiculo.get('year')}\n\n"""
    body += f"""<b>Servicios solicitados</b>\n{solicitud}"""
    body = body.replace('\n','<br>')
    email_text = """\
From: %s
To: %s
Subject: %s
Content-Type: text/html

%s
""" % (EMAIL_HOST_USER, ", ".join(to), subject, body)
    email_text = email_text.encode('utf-8')
    server = smtplib.SMTP_SSL('smtp.gmail.com')
    server.login(EMAIL_HOST_USER, EMAIL_HOST_PASSWORD)
    server.sendmail(EMAIL_HOST_USER, to, email_text)
    server.close()